# Developer: Jinglu Han
# mailbox: admin@de-manufacturing
import aiohttp
import io
import os
import tempfile
from docx import Document
from typing import Optional
import traceback
from app.services.logging import word_process_content_logger as logger

# 导入额外的处理库
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.warning("docx2txt库未安装，将跳过该处理方法")


async def extract_word_content_from_url(url: str) -> Optional[str]:
    """
    从指定URL下载Word文档并提取其中的文本内容

    Args:
        url: Word文档的下载地址

    Returns:
        str: 提取的文本内容，如果出错则返回None
    """
    temp_file_path = None
    try:
        # 创建异步HTTP会话，添加请求头模拟浏览器
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'application/msword, application/vnd.openxmlformats-officedocument.wordprocessingml.document, */*',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
        }

        async with aiohttp.ClientSession(headers=headers) as session:
            # 发送GET请求下载文件
            logger.info(f"正在下载文件: {url}")
            async with session.get(url) as response:
                if response.status != 200:
                    logger.error(f"下载失败，状态码: {response.status}")
                    return None

                # 读取响应内容
                content = await response.read()
                content_length = len(content)
                logger.info(f"下载完成，文件大小: {content_length} 字节")

                if content_length == 0:
                    logger.error("下载的文件为空")
                    return None

                # 使用临时文件保存下载的文档
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_file_path = temp_file.name
                    temp_file.write(content)

                logger.info(f"文件已保存到临时位置: {temp_file_path}")

                try:
                    # 使用python-docx解析文档
                    logger.info("开始解析文档...")
                    doc = Document(temp_file_path)
                    logger.info("文档解析成功")

                    # 提取所有段落的文本
                    text_content = []
                    for para in doc.paragraphs:
                        if para.text.strip():  # 忽略空段落
                            text_content.append(para.text)

                    # 提取表格中的文本
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():  # 忽略空单元格
                                    text_content.append(cell.text)

                    return "\n".join(text_content)
                except Exception as doc_error:
                    logger.debug(f"解析文档时出错: {str(doc_error)}")
                    logger.debug(f"详细错误信息: {traceback.format_exc()}")

                    # 尝试使用备用方法
                    try:
                        logger.info("尝试使用备用方法解析文档...")
                        # 检查文件是否为有效的ZIP文件（DOCX实际上是ZIP格式）
                        import zipfile
                        if zipfile.is_zipfile(temp_file_path):
                            with zipfile.ZipFile(temp_file_path) as zip_ref:
                                # 打印ZIP文件内容列表，帮助调试
                                logger.info(f"ZIP文件内容: {zip_ref.namelist()}")

                                # 尝试提取document.xml
                                if 'word/document.xml' in zip_ref.namelist():
                                    import xml.etree.ElementTree as ET
                                    with zip_ref.open('word/document.xml') as xml_file:
                                        tree = ET.parse(xml_file)
                                        root = tree.getroot()
                                        # 查找所有文本节点
                                        ns = {
                                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                                        }
                                        text_parts = []
                                        for t in root.findall('.//w:t', ns):
                                            if t.text:
                                                text_parts.append(t.text)
                                        return "\n".join(text_parts)
                        else:
                            logger.debug("文件不是有效的ZIP/DOCX格式")
                            return None
                    except Exception as backup_error:
                        logger.error(f"备用解析方法失败: {str(backup_error)}")
                        return None
                finally:
                    # 清理临时文件
                    if temp_file_path and os.path.exists(temp_file_path):
                        os.remove(temp_file_path)
                        logger.info(f"临时文件已删除: {temp_file_path}")

    except Exception as e:
        logger.error(f"处理Word文档时出错: {str(e)}")
        logger.error(f"详细错误信息: {traceback.format_exc()}")
        # 确保临时文件被删除
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            logger.info(f"临时文件已删除: {temp_file_path}")
        return None




async def extract_word_content_from_file(file_path: str) -> Optional[str]:
    """
    从本地Word文档文件提取文本内容

    Args:
        file_path: 本地Word文档文件的路径

    Returns:
        str: 提取的文本内容，如果出错则返回None
    """
    try:
        logger.info(f"开始解析本地文档: {file_path}")
        # 使用python-docx解析文档
        doc = Document(file_path)
        logger.info("文档解析成功")

        # 提取所有段落的文本
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():  # 忽略空段落
                text_content.append(para.text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # 忽略空单元格
                        text_content.append(cell.text)

        return "\n".join(text_content)
    except Exception as doc_error:
        logger.debug(f"解析文档时出错: {str(doc_error)}")
        logger.debug(f"详细错误信息: {traceback.format_exc()}")

        # 尝试使用备用方法
        try:
            logger.info("尝试使用备用方法解析文档...")
            # 检查文件是否为有效的ZIP文件（DOCX实际上是ZIP格式）
            import zipfile
            if zipfile.is_zipfile(file_path):
                with zipfile.ZipFile(file_path) as zip_ref:
                    # 打印ZIP文件内容列表，帮助调试
                    logger.info(f"ZIP文件内容: {zip_ref.namelist()}")

                    # 尝试提取document.xml
                    if 'word/document.xml' in zip_ref.namelist():
                        import xml.etree.ElementTree as ET
                        with zip_ref.open('word/document.xml') as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            # 查找所有文本节点
                            ns = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }
                            text_parts = []
                            for t in root.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)
                            return "\n".join(text_parts)
            else:
                logger.debug("文件不是有效的ZIP/DOCX格式")
                return None
        except Exception as backup_error:
            logger.error(f"备用解析方法失败: {str(backup_error)}")
            return None

    except Exception as e:
        logger.error(f"处理Word文档时出错: {str(e)}")
        logger.error(f"详细错误信息: {traceback.format_exc()}")
        return None


async def extract_word_content_from_bytes(file_content: bytes) -> Optional[str]:
    """
    从Word文档的字节流中提取文本内容，采用多重后备策略确保最大兼容性

    Args:
        file_content: Word文档的字节内容

    Returns:
        str: 提取的文本内容，如果所有方法都失败则返回错误信息
    """
    
    # 验证文件内容
    if not file_content or len(file_content) == 0:
        logger.error("文件内容为空或无效")
        return "错误：文件内容为空或无效"
    
    logger.info(f"开始处理Word文档，文件大小: {len(file_content)} 字节")
    
    # 策略1: 使用python-docx库
    try:
        logger.info("策略1: 使用python-docx解析...")
        file_like_object = io.BytesIO(file_content)
        
        doc = Document(file_like_object)
        logger.info("python-docx解析成功")

        # 提取所有段落的文本
        text_content = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_content.append(para.text.strip())

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_content.append(cell.text.strip())
        
        # 提取页眉页脚文本
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text and para.text.strip():
                        text_content.append(f"[页眉] {para.text.strip()}")
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text and para.text.strip():
                        text_content.append(f"[页脚] {para.text.strip()}")

        if text_content:
            result = "\n".join(text_content)
            logger.info(f"python-docx成功提取文本，长度: {len(result)} 字符")
            return result
        else:
            logger.warning("python-docx解析成功但未提取到任何文本内容")
    
    except Exception as e:
        logger.warning(f"python-docx解析失败: {str(e)}")
    
    # 策略2: 使用改进的XML解析方法
    try:
        logger.info("策略2: 使用改进的XML解析方法...")
        import zipfile
        import xml.etree.ElementTree as ET
        
        file_like_object = io.BytesIO(file_content)
        
        if not zipfile.is_zipfile(file_like_object):
            logger.error("文件不是有效的ZIP格式（.docx应该是ZIP格式）")
            return "错误：文件格式无效，.docx文件应该是ZIP格式"
        
        with zipfile.ZipFile(file_like_object) as zf:
            xml_parts = []
            
            # 首先尝试主要的文档内容
            if 'word/document.xml' in zf.namelist():
                try:
                    with zf.open('word/document.xml') as xml_file:
                        content = xml_file.read()
                        # 处理可能的编码问题
                        if content.startswith(b'\xef\xbb\xbf'):  # BOM
                            content = content[3:]
                        
                        tree = ET.fromstring(content)
                        
                        # 使用多个命名空间尝试
                        namespaces = [
                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'},
                            {'w': 'http://schemas.microsoft.com/office/word/2003/wordml'}
                        ]
                        
                        for ns in namespaces:
                            try:
                                # 提取文本节点
                                for t in tree.iterfind('.//w:t', ns):
                                    if t.text and t.text.strip():
                                        xml_parts.append(t.text.strip())
                                        
                                # 提取表格文本
                                for cell in tree.iterfind('.//w:tc', ns):
                                    cell_texts = []
                                    for t in cell.iterfind('.//w:t', ns):
                                        if t.text and t.text.strip():
                                            cell_texts.append(t.text.strip())
                                    if cell_texts:
                                        xml_parts.append(' '.join(cell_texts))
                                        
                                if xml_parts:
                                    break  # 如果找到了文本，就不尝试其他命名空间
                            except Exception:
                                continue
                                
                except ET.ParseError as pe:
                    logger.warning(f"XML解析错误: {pe}")
                except Exception as ee:
                    logger.warning(f"处理document.xml时出错: {ee}")
            
            # 如果主文档没有内容，尝试其他XML文件
            if not xml_parts:
                for filename in zf.namelist():
                    if filename.startswith('word/') and filename.endswith('.xml') and filename != 'word/document.xml':
                        try:
                            with zf.open(filename) as xml_file:
                                tree = ET.parse(xml_file)
                                root = tree.getroot()
                                # 提取任何包含文本的节点
                                for elem in root.iter():
                                    if elem.text and elem.text.strip() and len(elem.text.strip()) > 1:
                                        xml_parts.append(elem.text.strip())
                        except Exception:
                            continue
            
            if xml_parts:
                # 去重并过滤
                unique_parts = list(dict.fromkeys(xml_parts))  # 保持顺序的去重
                filtered_parts = [part for part in unique_parts if len(part) > 1]  # 过滤单字符
                
                if filtered_parts:
                    result = "\n".join(filtered_parts)
                    logger.info(f"XML方法成功提取文本，长度: {len(result)} 字符")
                    return result
    
    except Exception as e:
        logger.error(f"XML解析方法失败: {str(e)}")
        logger.error(f"详细错误: {traceback.format_exc()}")
    
    # 策略3: 使用docx2txt库（如果可用）
    if DOCX2TXT_AVAILABLE:
        try:
            logger.info("策略3: 使用docx2txt库...")
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file.flush()
                
                text_content = docx2txt.process(temp_file.name)
                
                # 清理临时文件
                try:
                    os.unlink(temp_file.name)
                except:
                    pass
                
                if text_content and text_content.strip():
                    result = text_content.strip()
                    logger.info(f"docx2txt成功提取文本，长度: {len(result)} 字符")
                    return result
                else:
                    logger.warning("docx2txt提取的内容为空")
        
        except Exception as e:
            logger.warning(f"docx2txt处理失败: {str(e)}")
    
    # 策略4: 尝试原始ZIP内容提取（最后的尝试）
    try:
        logger.info("策略4: 尝试原始内容提取...")
        file_like_object = io.BytesIO(file_content)
        
        if zipfile.is_zipfile(file_like_object):
            with zipfile.ZipFile(file_like_object) as zf:
                # 尝试读取所有文本文件
                text_parts = []
                for filename in zf.namelist():
                    if any(filename.endswith(ext) for ext in ['.xml', '.rels']):
                        try:
                            with zf.open(filename) as f:
                                content = f.read()
                                # 尝试解码为文本
                                try:
                                    text_content = content.decode('utf-8')
                                except UnicodeDecodeError:
                                    try:
                                        text_content = content.decode('gbk')
                                    except UnicodeDecodeError:
                                        continue
                                
                                # 使用正则表达式提取可能的文本内容
                                import re
                                # 查找XML文本节点内容
                                text_matches = re.findall(r'>([^<]{2,})<', text_content)
                                for match in text_matches:
                                    cleaned = match.strip()
                                    if len(cleaned) > 2 and not cleaned.isdigit():
                                        text_parts.append(cleaned)
                        except Exception:
                            continue
                
                if text_parts:
                    # 去重并排序
                    unique_parts = list(set(text_parts))
                    result = "\n".join(unique_parts)
                    logger.info(f"原始内容提取成功，长度: {len(result)} 字符")
                    return result
    
    except Exception as e:
        logger.error(f"原始内容提取失败: {str(e)}")
    
    # 所有策略都失败了
    error_msg = "无法从Word文档中提取文本内容。可能的原因：\n1. 文件已损坏\n2. 文件格式不受支持\n3. 文件受密码保护\n4. 文件内容为纯图片\n\n建议：请尝试用Microsoft Word打开并重新保存该文件，或转换为PDF格式后重试。"
    logger.error("所有文本提取策略均失败")
    return error_msg
    
# 使用示例
# async def main():
#     url = "https://x/img/2025/05/07/14/6e/6e811b83-8f22-43d2-894b-8fc236ff971f.docx"
#     print(f"开始处理URL: {url}")
#     content = await extract_word_content_from_url(url)
#     if content:
#         print("提取的内容:")
#         print(content)
#     else:
#         print("无法提取内容")


# # 如果直接运行此文件
# if __name__ == "__main__":
#     import asyncio

#     asyncio.run(main())